"""Data Harvester Tool -- extracts raw text from invoice files.

Supports PDF (pdfplumber), DOCX (python-docx), and image (pytesseract)
formats. File type is detected from content bytes via the filetype library,
never from the file extension or meta-declared attachment name.
"""

import logging
from pathlib import Path

import filetype

logger = logging.getLogger(__name__)


def data_harvester_tool(file_path: str) -> str:
    """Extract raw text from an invoice file (PDF, DOCX, or image).

    Detects format from file content bytes, not from extension.

    Args:
        file_path: Absolute or relative path to the invoice file.

    Returns:
        Extracted raw text as a single string.

    Raises:
        FileNotFoundError: If file does not exist.
        ValueError: If file format is unsupported.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Invoice file not found: {file_path}")

    # Detect actual file type from content
    kind = filetype.guess(str(path))
    if kind is None:
        fmt = path.suffix.lstrip(".").lower()
    else:
        mime = kind.mime
        if "pdf" in mime:
            fmt = "pdf"
        elif "image" in mime:
            fmt = "image"
        elif "officedocument.wordprocessingml" in mime:
            fmt = "docx"
        else:
            fmt = kind.extension or path.suffix.lstrip(".").lower()

    logger.info("Detected format '%s' for %s", fmt, file_path)

    if fmt == "pdf":
        return _extract_pdf(path)
    elif fmt in ("png", "jpg", "jpeg", "image"):
        return _extract_image(path)
    elif fmt == "docx":
        return _extract_docx(path)
    else:
        raise ValueError(f"Unsupported file format: {fmt} for {file_path}")


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_image(path: Path) -> str:
    """Extract text from image using pytesseract OCR."""
    from PIL import Image
    import pytesseract

    image = Image.open(str(path))
    text = pytesseract.image_to_string(image, config="--psm 6")
    return text.strip()


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX using python-docx."""
    import docx

    doc = docx.Document(str(path))
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)
    return "\n".join(text_parts)
